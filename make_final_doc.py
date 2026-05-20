"""Generate the final script doc from finalized.json -> docx -> Google Doc.

Built with python-docx for direct control over font + spacing.
Body: Georgia 14pt, line spacing 1.6, generous paragraph spacing.
Title + section headings: Helvetica Neue, bold, tight letter spacing.
"""
import json, subprocess, os, datetime, time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING

fz = json.load(open('finalized.json'))
all_entries = fz['finalized']
entries = [e for e in all_entries if not e.get('rework_note')]
TITLES = {s['slide_id']: s.get('video_title', '') for s in json.load(open('slides.json'))}

doc = Document()

# Generous margins
for section in doc.sections:
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.2)
    section.left_margin = Inches(1.4)
    section.right_margin = Inches(1.4)

BODY_FONT = 'Georgia'
HEAD_FONT = 'Helvetica Neue'
INK = RGBColor(0x1a, 0x1a, 0x1a)
MUTE = RGBColor(0x6a, 0x6a, 0x6a)

def set_run(run, font_name, size_pt, bold=False, color=INK):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = color

def add_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(36)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, HEAD_FONT, 26, bold=True, color=INK)

def add_section_head(num, title):
    # Numbered heading with the short title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_with_next = True
    num_run = p.add_run(f'{num:02d}  ')
    set_run(num_run, HEAD_FONT, 13, bold=False, color=MUTE)
    title_run = p.add_run(title)
    set_run(title_run, HEAD_FONT, 17, bold=True, color=INK)

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run(run, BODY_FONT, 13, bold=False, color=INK)

add_title('SEA Brand Channel Batch 1 - Final Scripts')

for i, e in enumerate(entries, 1):
    title = TITLES.get(e['slide_id'], '')
    add_section_head(i, title if title else '')
    # Final VO might contain multiple sentences; split on double newlines if present, otherwise single paragraph
    paragraphs = [p.strip() for p in e['final_vo'].split('\n\n') if p.strip()] or [e['final_vo']]
    for para in paragraphs:
        add_body(para)

OUT_DOCX = 'SEA_Batch1_Final_Scripts.docx'
doc.save(OUT_DOCX)
print(f'Built docx: {len(entries)} scripts')

# --- Upload + convert to Google Doc ---
import urllib.request
cfg = json.loads(subprocess.run(['/opt/homebrew/bin/rclone','config','dump'],capture_output=True,text=True).stdout)
tok = json.loads(cfg['gdrive']['token'])['access_token']
DEST = "gdrive:Video Scripts/SEA Brand Channel/Batch 1 FINAL/"
ts = str(int(time.time()))
tmp = f'SEA_Batch1_tmp_{ts}.docx'
subprocess.run(['cp', OUT_DOCX, tmp], check=True)
subprocess.run(['/opt/homebrew/bin/rclone', 'copy', tmp, DEST, '--no-traverse'], check=True)
docx_id = json.loads(subprocess.run(['/opt/homebrew/bin/rclone','lsjson',DEST+tmp],capture_output=True,text=True).stdout)[0]['ID']
body = json.dumps({'name': 'SEA_Batch1_Final_Scripts', 'mimeType': 'application/vnd.google-apps.document'}).encode()
res = json.load(urllib.request.urlopen(urllib.request.Request(
    f'https://www.googleapis.com/drive/v3/files/{docx_id}/copy', data=body,
    headers={'Authorization': f'Bearer {tok}', 'Content-Type': 'application/json'}, method='POST')))
new_doc = res['id']
old = open('.final_doc_id').read().strip() if os.path.exists('.final_doc_id') else ''
for fid in [x for x in (docx_id, old) if x]:
    try: urllib.request.urlopen(urllib.request.Request(f'https://www.googleapis.com/drive/v3/files/{fid}', headers={'Authorization': f'Bearer {tok}'}, method='DELETE'))
    except Exception: pass
open('.final_doc_id', 'w').write(new_doc)
url = f'https://docs.google.com/document/d/{new_doc}/edit'
open('/tmp/sea_final_doc_url.txt', 'w').write(url)
subprocess.run(['rm', '-f', tmp])
print(f'Doc: {url}')
